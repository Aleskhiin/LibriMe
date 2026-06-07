import asyncio
import os
import subprocess
import tempfile
from pathlib import Path

from docx import Document

from .BaseDocumentReader import BaseDocumentReader


class ODTDocumentReader(BaseDocumentReader):
    """
    Reader implementation for ODT documents.

    The ODT file is converted to DOCX through LibreOffice and then
    read with python-docx.
    """

    def __init__(self):
        """Initializes the ODT reader state."""
        super().__init__()
        self.paragraphs: list[str] = []

    def _convert_odt_to_docx(self, odt_path: str) -> str:
        """
        Converts an .odt file to .docx using LibreOffice.

        Args:
            odt_path: Path to the ODT file.

        Returns:
            Path to the converted DOCX file.
        """

        temp_dir = tempfile.mkdtemp()

        subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                temp_dir,
                odt_path
            ],
            check=True
        )

        docx_file = Path(temp_dir) / f"{Path(odt_path).stem}.docx"

        if not docx_file.exists():
            raise RuntimeError(f"Conversion failed for {odt_path}")

        return str(docx_file)

    def _load_document(self) -> None:
        """
        Loads the converted ODT document and extracts paragraphs.
        """

        if not self.file_path:
            raise ValueError("ODT file path not configured.")

        converted_file = self._convert_odt_to_docx(self.file_path)
        document = Document(converted_file)

        self.paragraphs = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]

    def read_document(self) -> str:
        """
        Returns the ODT document text.
        """

        return "\n\n".join(self.paragraphs)

    def read_paragraphs(self) -> list[str]:
        """
        Returns extracted ODT paragraphs.
        """

        return self.paragraphs


async def main():
    """
    Tests the ODTDocumentReader directly.
    """

    base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    file_path = os.path.join(base_dir, "Testfiles", "example.odt")

    reader = ODTDocumentReader()
    reader.configure(file_path=file_path, reader_mode="document")

    result = await reader.process()

    print("\n=== ODT DOCUMENT RESULT ===")
    print(result)


if __name__ == "__main__":
    asyncio.run(main())