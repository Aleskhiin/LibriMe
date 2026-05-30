import asyncio
import os
import subprocess
import tempfile
from pathlib import Path

from docx import Document

from .BaseDocumentReader import BaseDocumentReader


class WordDocumentReader(BaseDocumentReader):
    """
    Reader implementation for Word documents.

    Supports:
    - .docx directly through python-docx
    - .doc by converting it to .docx through LibreOffice
    """

    def __init__(self):
        """Initializes the Word reader state."""
        super().__init__()
        self.paragraphs: list[str] = []

    def _convert_doc_to_docx(self, doc_path: str) -> str:
        """
        Converts a legacy .doc file to .docx using LibreOffice.

        Args:
            doc_path: Path to the .doc file.

        Returns:
            Path to the converted .docx file.
        """

        temp_dir = tempfile.mkdtemp()

        # Run LibreOffice in headless mode for Linux server environments
        subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                temp_dir,
                doc_path
            ],
            check=True
        )

        docx_file = Path(temp_dir) / f"{Path(doc_path).stem}.docx"

        if not docx_file.exists():
            raise RuntimeError(f"Conversion failed for {doc_path}")

        return str(docx_file)

    def _load_document(self) -> None:
        """
        Loads the Word document and extracts all non-empty paragraphs.
        """

        if not self.file_path:
            raise ValueError("Word file path not configured.")

        path = Path(self.file_path)
        suffix = path.suffix.lower()

        # Convert legacy .doc files before reading
        if suffix == ".doc":
            file_to_open = self._convert_doc_to_docx(self.file_path)

        # Read modern .docx files directly
        elif suffix == ".docx":
            file_to_open = self.file_path

        else:
            raise ValueError(f"Unsupported Word format: {suffix}")

        document = Document(file_to_open)

        self.paragraphs = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]

    def read_document(self) -> str:
        """
        Returns the Word document text.
        """

        return "\n\n".join(self.paragraphs)

    def read_paragraphs(self) -> list[str]:
        """
        Returns extracted Word paragraphs.
        """

        return self.paragraphs


async def main():
    """
    Tests the WordDocumentReader directly.
    """

    base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    testfiles_dir = os.path.join(base_dir, "Testfiles")

    test_files = [
        "example.docx",
        "example.doc"
    ]

    for file_name in test_files:
        file_path = os.path.join(testfiles_dir, file_name)

        if not os.path.exists(file_path):
            print(f"\nSkipping {file_name}: file does not exist.")
            continue

        reader = WordDocumentReader()
        reader.configure(file_path=file_path, reader_mode="document")

        result = await reader.process()

        print(f"\n=== WORD RESULT: {file_name} ===")
        print(result)


if __name__ == "__main__":
    asyncio.run(main())